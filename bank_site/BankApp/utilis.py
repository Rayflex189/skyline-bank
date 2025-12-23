from .decorators import *
from .forms import *
from .models import *
from django.core.signing import TimestampSigner, BadSignature, SignatureExpired


signer = TimestampSigner()
from django.core.mail import EmailMessage

from docx import Document
from docx.shared import Pt
from django.http import HttpResponse


def generate_loan_pdf(loan):
    doc = Document()
    doc.add_heading("Loan Approval Certificate", level=1)

    doc.add_paragraph(f"Name: {loan.full_name}")
    doc.add_paragraph(f"Loan Amount: ₦{loan.loan_amount}")
    doc.add_paragraph(f"Processing Fee: ₦{loan.processing_fee}")
    doc.add_paragraph(f"Total Due: ₦{loan.total_amount_due}")

    filename = f"Loan_{loan.id}_Approval.docx"
    file_path = f"/tmp/{filename}"
    doc.save(file_path)

    return file_path


def email_pdf(user_email, pdf_path):
    email = EmailMessage(
        "Loan Approved – Your Approval Letter",
        "Dear customer,\n\nYour loan request has been approved. Please find the approval letter attached.",
        settings.EMAIL_HOST_USER,
        [user_email],
    )
    email.attach_file(pdf_path)
    email.send()


def calculate_interest(amount, duration):
    months = int(duration.split()[0])

    if months <= 3:
        rate = 5
    elif months <= 6:
        rate = 10
    else:
        rate = 15

    interest_amount = (amount * rate) / 100
    total = amount + interest_amount

    return rate, total


# utils.py
def validate_otp(input_otp, user_profile):
    # Compare the input OTP with the OTP stored in the user profile
    return input_otp == user_profile.otp_code

def validate_imf(input_imf, user_profile):
    # Compare the input OTP with the OTP stored in the user profile
    return input_imf == user_profile.imf_code

def validate_aml(input_aml, user_profile):
    # Compare the input OTP with the OTP stored in the user profile
    return input_aml == user_profile.aml_code

def validate_tac(input_tac, user_profile):
    # Compare the input OTP with the OTP stored in the user profile
    return input_tac == user_profile.tac_code

def validate_vat(input_vat, user_profile):
    # Compare the input OTP with the OTP stored in the user profile
    return input_vat == user_profile.vat_code

def validate_linking_code(input_linking_code, user_profile):
    # Compare the input OTP with the OTP stored in the user profile
    return input_linking_code == user_profile.linking_code