# At the top of views.py, add this import:
from datetime import date, datetime
import datetime as dt  # Alternative if you need both
# Django Core Imports
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.db.models import Count, Sum, Avg, Q
import re
from django.contrib.auth import get_user_model
from django.db import transaction
from django.conf import settings
from django.utils.http import urlsafe_base64_decode
from django.utils.encoding import force_str
from django.contrib.auth.tokens import default_token_generator
from django.contrib import messages
from django.core.exceptions import ValidationError, ObjectDoesNotExist
from django.core.signing import TimestampSigner, BadSignature, SignatureExpired
from django.core.mail import send_mail
from django.urls import reverse
from django.db.models import Q
import datetime
from decimal import Decimal

# Utility Imports
from datetime import timedelta, datetime
from django.utils import timezone
from datetime import timedelta
from django.contrib.auth import authenticate, login, logout

# Project Imports
from .decorators import *
from .forms import *
from .models import *
from .utilis import *  # If still required (consider limiting *)
from BankApp.decorators import unauthenticated_user
from BankApp.models import UserProfile
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import InvestmentPlan, UserInvestment, UserProfile
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib.auth.models import User
from django.db.models import Count, Q
import datetime

from django.contrib.admin.views.decorators import staff_member_required

from django.http import HttpResponse
from docx import Document


User = get_user_model()
signer = TimestampSigner()


def submit_loan(request):
    if request.method == "POST":
        form = LoanForm(request.POST)

        if form.is_valid():
            loan = form.save(commit=False)
            loan.user = request.user

            # calculate interest
            rate, total = calculate_interest(loan.loan_amount, loan.loan_duration)
            loan.interest_rate = rate
            loan.total_amount_due = total

            loan.save()

            # generate PDF
            pdf_path = generate_loan_pdf(loan)

            # email PDF
            email_pdf(request.user.email, pdf_path)

            messages.success(request, "Loan submitted successfully! Approval sent to email.")
            return redirect("dashboard")
    else:
        form = LoanForm()

    return render(request, "BankApp/submit_loans.html", {"form": form})


def send_kyc_email(request, user_id):
    user = get_object_or_404(User, id=user_id)

    # Attach the user's latest KYC if you want
    try:
        kyc = KYC.objects.get(user=user)
    except KYC.DoesNotExist:
        kyc = None

    subject = "Your KYC Verification Status"
    message = (
        f"Dear {user.first_name},\n\n"
        "Thank you for submitting your KYC details. "
        "Your verification is being reviewed.\n\n"
        "Best regards,\nYour Bank Team"
    )

    email = EmailMessage(subject, message, settings.EMAIL_HOST_USER, [user.email])
    email.send()

    messages.success(request, "KYC email sent successfully.")
    return redirect("dashboard")  # or anywhere you want



def download_kyc_pdf(request, user_id):
    user_profile = get_object_or_404(UserProfile, user_id=user_id)

    # Create a DOCX file
    doc = Document()
    doc.add_heading("KYC Details", level=1)

    doc.add_paragraph(f"Full Name: {user_profile.full_name}")
    doc.add_paragraph(f"Email: {user_profile.user.email}")
    doc.add_paragraph(f"Phone: {user_profile.phone}")
    doc.add_paragraph(f"Address: {user_profile.address}")
    doc.add_paragraph(f"Occupation: {user_profile.occupation}")
    doc.add_paragraph(f"City: {user_profile.city}")
    doc.add_paragraph(f"ZIP Code: {user_profile.zip_code}")

    # Save to buffer
    file_name = f"KYC_{user_profile.user.email}.docx"
    file_path = f"/tmp/{file_name}"
    doc.save(file_path)

    with open(file_path, "rb") as file:
        response = HttpResponse(file.read(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response["Content-Disposition"] = f'attachment; filename="{file_name}"'
        return response



@staff_member_required
def manage_loans(request):
    """
    Main loan management view for staff members.
    """
    # Get all loans ordered by submission date (most recent first)
    loans = Loan.objects.all().order_by('-submitted_at')
    
    # Get filter parameters from GET request
    status_filter = request.GET.get('status', '')
    search_query = request.GET.get('q', '')
    
    # Apply filters
    if status_filter:
        loans = loans.filter(status__iexact=status_filter)
    
    if search_query:
        loans = loans.filter(
            Q(user__username__icontains=search_query) |
            Q(user__email__icontains=search_query) |
            Q(loan_type__icontains=search_query) |
            Q(purpose__icontains=search_query) |
            Q(user__first_name__icontains=search_query) |
            Q(user__last_name__icontains=search_query)
        )
    
    # Get statistics for the dashboard
    total_loans = Loan.objects.count()
    pending_loans = Loan.objects.filter(status__iexact='Pending').count()
    approved_loans = Loan.objects.filter(status__iexact='Approved').count()
    rejected_loans = Loan.objects.filter(status__iexact='Rejected').count()
    
    # Add calculated fields to each loan for the template
    for loan in loans:
        # Calculate processing fee (5% of loan amount) - convert to Decimal first
        try:
            # Convert amount to Decimal if it isn't already
            amount_decimal = Decimal(str(loan.amount))
            loan.processing_fee = amount_decimal * Decimal('0.05')
            
            # Calculate total amount due (principal + processing fee)
            loan.total_amount_due = amount_decimal + loan.processing_fee
            
            # Format the values for display
            loan.processing_fee_display = f"${loan.processing_fee:,.2f}"
            loan.total_amount_due_display = f"${loan.total_amount_due:,.2f}"
            loan.amount_display = f"${loan.amount:,.2f}"
            
        except Exception as e:
            # Fallback values if calculation fails
            loan.processing_fee = Decimal('0.00')
            loan.total_amount_due = loan.amount
            loan.processing_fee_display = "$0.00"
            loan.total_amount_due_display = f"${loan.amount:,.2f}"
            loan.amount_display = f"${loan.amount:,.2f}"
        
        # Format user full name if available
        if loan.user:
            loan.full_name = f"{loan.user.first_name} {loan.user.last_name}".strip()
            if not loan.full_name:
                loan.full_name = loan.user.username
            loan.email = loan.user.email
        else:
            loan.full_name = "Unknown User"
            loan.email = "N/A"
        
        # For loan purpose display
        loan.loan_purpose = loan.purpose if loan.purpose else loan.loan_type
    
    # Handle bulk actions
    if request.method == 'POST' and 'action' in request.POST:
        loan_ids = request.POST.getlist('loan_ids')
        action = request.POST.get('action')
        
        if loan_ids and action:
            loans_to_process = Loan.objects.filter(id__in=loan_ids)
            
            if action == 'approve':
                loans_to_process.update(status='Approved')
                messages.success(request, f'{len(loan_ids)} loan(s) approved successfully.')
            elif action == 'reject':
                loans_to_process.update(status='Rejected')
                messages.success(request, f'{len(loan_ids)} loan(s) rejected.')
            
            return redirect('manage_loans')
    
    context = {
        'loans': loans,
        'total_count': total_loans,
        'pending_count': pending_loans,
        'approved_count': approved_loans,
        'rejected_count': rejected_loans,
        'status_filter': status_filter,
        'search_query': search_query,
        'status_choices': ['Pending', 'Approved', 'Rejected'],
    }
    
    return render(request, 'BankApp/manage_loans.html', context)


@staff_member_required
def approve_loan(request, loan_id):
    """
    Approve individual loan and send notification email.
    """
    try:
        loan = get_object_or_404(Loan, id=loan_id)
        
        # Check if already approved
        if loan.status == 'Approved':
            messages.warning(request, f"Loan #{loan.id} is already approved.")
            return redirect('manage_loans')
        
        # Update status and set reviewed timestamp
        loan.status = 'Approved'
        loan.reviewed_at = datetime.datetime.now()
        loan.save()
        
        # Calculate amounts - using Decimal for proper calculation
        amount_decimal = Decimal(str(loan.amount))
        processing_fee = amount_decimal * Decimal('0.05')
        total_due = amount_decimal + processing_fee
        
        # Get user email
        user_email = loan.user.email if loan.user else None
        
        # Send approval email if user has email
        if user_email:
            try:
                send_mail(
                    subject="Loan Approved - SkyBridge Finance",
                    message=(
                        f"Hello {loan.user.get_full_name() or loan.user.username},\n\n"
                        f"Congratulations! Your loan application (ID: #{loan.id}) has been approved.\n\n"
                        f"Loan Details:\n"
                        f"Amount: ${loan.amount:,.2f}\n"
                        f"Loan Type: {loan.loan_type}\n"
                        f"Duration: {loan.duration} months\n"
                        f"Interest Rate: {loan.interest}%\n"
                        f"Processing Fee (5%): ${processing_fee:,.2f}\n"
                        f"Total Amount Due: ${total_due:,.2f}\n"
                        f"Purpose: {loan.purpose if loan.purpose else loan.loan_type}\n\n"
                        f"Next Steps:\n"
                        f"1. Please login to your account\n"
                        f"2. Review the loan terms\n"
                        f"3. Accept the loan agreement\n\n"
                        f"You will receive the funds within 2-3 business days.\n\n"
                        "If you have any questions, please contact our support team.\n\n"
                        "Best regards,\n"
                        "SkyBridge Finance Team\n"
                        "support@skybridgefinance.online"
                    ),
                    from_email=settings.DEFAULT_FROM_EMAIL,
                    recipient_list=[user_email],
                    fail_silently=False,
                )
                email_sent = True
            except Exception as e:
                email_sent = False
                messages.warning(request, f"Loan approved but email failed to send: {str(e)}")
        else:
            email_sent = False
            messages.warning(request, f"Loan approved but no email found for user.")
        
        # Success message
        if email_sent:
            messages.success(request, f"Loan #{loan.id} approved and notification email sent.")
        else:
            messages.success(request, f"Loan #{loan.id} approved.")
            
    except Loan.DoesNotExist:
        messages.error(request, "Loan not found.")
    except Exception as e:
        messages.error(request, f"Error approving loan: {str(e)}")
    
    return redirect('manage_loans')


@staff_member_required
def reject_loan(request, loan_id):
    """
    Reject individual loan and send notification email.
    """
    try:
        loan = get_object_or_404(Loan, id=loan_id)
        
        # Check if already rejected
        if loan.status == 'Rejected':
            messages.warning(request, f"Loan #{loan.id} is already rejected.")
            return redirect('manage_loans')
        
        # Update status and set reviewed timestamp
        loan.status = 'Rejected'
        loan.reviewed_at = datetime.datetime.now()
        loan.save()
        
        # Get user email
        user_email = loan.user.email if loan.user else None
        
        # Get rejection reason from form if available
        rejection_reason = request.POST.get('rejection_reason', 
                                          "Your application did not meet our current lending criteria.")
        
        # Send rejection email if user has email
        if user_email:
            try:
                send_mail(
                    subject="Loan Application Update - SkyBridge Finance",
                    message=(
                        f"Hello {loan.user.get_full_name() or loan.user.username},\n\n"
                        f"We regret to inform you that your loan application (ID: #{loan.id}) "
                        f"for ${loan.amount:,.2f} has not been approved at this time.\n\n"
                        f"Reason: {rejection_reason}\n\n"
                        f"You may reapply in 30 days or contact our support team for more information.\n\n"
                        f"Application Details:\n"
                        f"Amount Requested: ${loan.amount:,.2f}\n"
                        f"Loan Type: {loan.loan_type}\n"
                        f"Purpose: {loan.purpose if loan.purpose else loan.loan_type}\n\n"
                        f"Thank you for considering SkyBridge Finance for your lending needs.\n\n"
                        "Best regards,\n"
                        "SkyBridge Finance Team\n"
                        "support@skybridgefinance.online"
                    ),
                    from_email=settings.DEFAULT_FROM_EMAIL,
                    recipient_list=[user_email],
                    fail_silently=False,
                )
                email_sent = True
            except Exception as e:
                email_sent = False
                messages.warning(request, f"Loan rejected but email failed to send: {str(e)}")
        else:
            email_sent = False
            messages.warning(request, f"Loan rejected but no email found for user.")
        
        # Success message
        if email_sent:
            messages.error(request, f"Loan #{loan.id} rejected and notification email sent.")
        else:
            messages.error(request, f"Loan #{loan.id} rejected.")
            
    except Loan.DoesNotExist:
        messages.error(request, "Loan not found.")
    except Exception as e:
        messages.error(request, f"Error rejecting loan: {str(e)}")
    
    return redirect('manage_loans')

@staff_member_required
def view_loan_details(request, loan_id):
    """
    View detailed information about a specific loan.
    """
    loan = get_object_or_404(Loan, id=loan_id)
    
    # Calculate additional information
    processing_fee = loan.loan_amount * 0.05
    total_due = loan.loan_amount + processing_fee
    
    context = {
        'loan': loan,
        'processing_fee': processing_fee,
        'total_due': total_due,
    }
    
    return render(request, 'BankApp/loan_details.html', context)



# Helper function to calculate interest (updated for new fields)
def calculate_interest(amount, duration, loan_type=None, annual_income=None):
    """
    Calculate interest rate based on loan parameters.
    Preserving original logic with enhancements for new fields.
    """
    # Base interest rate (preserving original logic)
    base_rate = 5.0  # 5% base
    
    # Adjust based on loan type if provided
    if loan_type:
        loan_type_lower = loan_type.lower()
        if 'mortgage' in loan_type_lower or loan_type == 'mortgage':
            base_rate = 3.5
        elif 'auto' in loan_type_lower or loan_type == 'auto':
            base_rate = 4.5
        elif 'education' in loan_type_lower or loan_type == 'education':
            base_rate = 4.0
        elif 'emergency' in loan_type_lower or loan_type in ['payday', 'emergency']:
            base_rate = 10.0
        elif 'business' in loan_type_lower or loan_type == 'business':
            base_rate = 6.5
        elif 'personal' in loan_type_lower or loan_type == 'personal':
            base_rate = 7.5
    
    # Adjust based on duration (longer terms might have higher rates)
    if duration > 60:  # More than 5 years
        base_rate += 0.5
    
    # Adjust based on income if provided
    if annual_income and amount:
        try:
            income_ratio = float(amount) / float(annual_income)
            if income_ratio > 2:
                base_rate += 1.0
            elif income_ratio < 0.5:
                base_rate -= 0.5
        except (ValueError, TypeError):
            pass  # Keep base rate if calculation fails
    
    # Calculate total payable (preserving original calculation method)
    # Using simple interest for backward compatibility
    total = float(amount) * (1 + (base_rate / 100) * (duration / 12))
    
    return base_rate, round(total, 2)

@staff_member_required
def admin_dashboard(request):
    """
    Comprehensive admin dashboard with statistics and analytics
    Updated for current Loan model structure
    """
    # User statistics
    total_users = UserProfile.objects.count()
    active_users = UserProfile.objects.filter(user__is_active=True).count()
    
    # Loan statistics - using proper status values from your model
    total_loans = Loan.objects.count()
    
    # Using exact status values (case-sensitive as stored in database)
    pending_loans = Loan.objects.filter(status='Pending').count()
    approved_loans = Loan.objects.filter(status='Approved').count()
    rejected_loans = Loan.objects.filter(status='Rejected').count()
    
    # Today's loans
    today = datetime.date.today()
    today_loans = Loan.objects.filter(submitted_at__date=today).count()
    
    # This week's loans (last 7 days)
    week_ago = datetime.datetime.now() - datetime.timedelta(days=7)
    this_week_loans = Loan.objects.filter(submitted_at__gte=week_ago).count()
    
    # This month's loans (last 30 days)
    month_ago = datetime.datetime.now() - datetime.timedelta(days=30)
    this_month_loans = Loan.objects.filter(submitted_at__gte=month_ago).count()
    
    # Financial statistics
    total_loan_amount = Loan.objects.aggregate(total=Sum('amount'))['total'] or Decimal('0')
    total_pending_amount = Loan.objects.filter(status='Pending').aggregate(total=Sum('amount'))['total'] or Decimal('0')
    total_approved_amount = Loan.objects.filter(status='Approved').aggregate(total=Sum('amount'))['total'] or Decimal('0')
    total_rejected_amount = Loan.objects.filter(status='Rejected').aggregate(total=Sum('amount'))['total'] or Decimal('0')
    
    # Average loan amount
    avg_loan_amount = Loan.objects.aggregate(avg=Avg('amount'))['avg'] or Decimal('0')
    
    # Loan type distribution
    loan_type_stats = Loan.objects.values('loan_type').annotate(
        count=Count('id'),
        total_amount=Sum('amount'),
        avg_amount=Avg('amount')
    ).order_by('-total_amount')
    
    # Status distribution for charts
    status_distribution = [
        {'status': 'Pending', 'count': pending_loans, 'color': '#f59e0b'},
        {'status': 'Approved', 'count': approved_loans, 'color': '#10b981'},
        {'status': 'Rejected', 'count': rejected_loans, 'color': '#ef4444'},
    ]
    
    # Duration statistics
    duration_stats = Loan.objects.values('duration').annotate(
        count=Count('id'),
        avg_amount=Avg('amount')
    ).order_by('duration')
    
    # Recent pending loans (for quick action)
    recent_pending_loans = Loan.objects.filter(
        status='Pending'
    ).select_related('user').order_by('-submitted_at')[:5]
    
    # Add processing fee calculation for display
    for loan in recent_pending_loans:
        try:
            amount_decimal = Decimal(str(loan.amount))
            loan.processing_fee = amount_decimal * Decimal('0.05')
            loan.total_due = amount_decimal + loan.processing_fee
        except:
            loan.processing_fee = Decimal('0')
            loan.total_due = loan.amount
    
    # Recent approved loans (for monitoring)
    recent_approved_loans = Loan.objects.filter(
        status='Approved'
    ).select_related('user').order_by('-reviewed_at')[:5]
    
    # Loan approval rate
    total_processed = approved_loans + rejected_loans
    approval_rate = (approved_loans / total_processed * 100) if total_processed > 0 else 0
    
    # Monthly trend data (last 6 months)
    monthly_trends = []
    for i in range(6, -1, -1):
        month_start = datetime.datetime.now().replace(
            day=1, hour=0, minute=0, second=0, microsecond=0
        ) - datetime.timedelta(days=30*i)
        month_end = month_start + datetime.timedelta(days=30)
        
        month_loans = Loan.objects.filter(
            submitted_at__gte=month_start,
            submitted_at__lt=month_end
        ).aggregate(
            count=Count('id'),
            amount=Sum('amount')
        )
        
        monthly_trends.append({
            'month': month_start.strftime('%b %Y'),
            'count': month_loans['count'] or 0,
            'amount': month_loans['amount'] or Decimal('0'),
        })
    
    # Top 5 loan amounts
    top_loans = Loan.objects.select_related('user').order_by('-amount')[:5]
    
    # Loan purposes frequency
    purpose_stats = Loan.objects.exclude(purpose__isnull=True).exclude(purpose='').values(
        'purpose'
    ).annotate(
        count=Count('id')
    ).order_by('-count')[:10]
    
    context = {
        # User stats
        'total_users': total_users,
        'active_users': active_users,
        
        # Loan count stats
        'total_loans': total_loans,
        'pending_loans': pending_loans,
        'approved_loans': approved_loans,
        'rejected_loans': rejected_loans,
        'today_loans': today_loans,
        'this_week_loans': this_week_loans,
        'this_month_loans': this_month_loans,
        
        # Financial stats
        'total_loan_amount': total_loan_amount,
        'total_pending_amount': total_pending_amount,
        'total_approved_amount': total_approved_amount,
        'total_rejected_amount': total_rejected_amount,
        'avg_loan_amount': avg_loan_amount,
        'approval_rate': round(approval_rate, 1),
        
        # Analytics data
        'loan_type_stats': loan_type_stats,
        'status_distribution': status_distribution,
        'duration_stats': duration_stats,
        'purpose_stats': purpose_stats,
        'monthly_trends': monthly_trends,
        
        # Recent data for quick access
        'recent_pending_loans': recent_pending_loans,
        'recent_approved_loans': recent_approved_loans,
        'top_loans': top_loans,
        
        # Formatting helpers
        'now': datetime.datetime.now(),
    }
    
    return render(request, 'BankApp/admin_dashboard.html', context)

# Add this helper function for debugging
def check_staff_status(request):
    """
    Debug view to check if user has staff permissions.
    Accessible at /staff/check/
    """
    if request.user.is_authenticated:
        user_info = {
            'username': request.user.username,
            'email': request.user.email,
            'is_staff': request.user.is_staff,
            'is_superuser': request.user.is_superuser,
            'is_authenticated': request.user.is_authenticated,
        }
        messages.info(request, f"User info: {user_info}")
    
    return redirect('manage_loans')

@login_required
def apply_loan(request):
    """
    Apply for a loan with new form fields.
    Preserving original functionality while adding new fields.
    """
    try:
        user_profile = UserProfile.objects.get(user=request.user)
    except UserProfile.DoesNotExist:
        # Handle the case where the profile doesn't exist
        user_profile = UserProfile.objects.create(user=request.user)
    
    if request.method == "POST":
        form = LoanForm(request.POST)
        if form.is_valid():
            # Get all cleaned data
            amount = form.cleaned_data['amount']
            loan_type = form.cleaned_data['loan_type']
            duration = form.cleaned_data['duration']
            
            # Get new fields if they exist in the form
            purpose = form.cleaned_data.get('purpose', 'other')
            employment_status = form.cleaned_data.get('employment_status', 'employed')
            annual_income = form.cleaned_data.get('annual_income', 0)
            repayment_frequency = form.cleaned_data.get('repayment_frequency', 'monthly')
            collateral = form.cleaned_data.get('collateral', '')
            requested_date = form.cleaned_data.get('requested_date', date.today())
            notes = form.cleaned_data.get('notes', '')
            
            # Calculate interest (updated to accept new parameters)
            interest_rate, total = calculate_interest(
                float(amount), 
                duration, 
                loan_type, 
                annual_income
            )
            
            # Store in session for review
            loan_data = {
                'amount': float(amount),
                'loan_type': loan_type,
                'duration': duration,
                'interest': interest_rate,
                'total': float(total)
            }
            
            # Add new fields to session data if they exist
            if 'purpose' in form.cleaned_data:
                loan_data['purpose'] = purpose
            if 'employment_status' in form.cleaned_data:
                loan_data['employment_status'] = employment_status
            if 'annual_income' in form.cleaned_data:
                loan_data['annual_income'] = float(annual_income) if annual_income else 0
            if 'repayment_frequency' in form.cleaned_data:
                loan_data['repayment_frequency'] = repayment_frequency
            if 'collateral' in form.cleaned_data:
                loan_data['collateral'] = collateral
            if 'requested_date' in form.cleaned_data:
                loan_data['requested_date'] = requested_date.isoformat()
            if 'notes' in form.cleaned_data:
                loan_data['notes'] = notes
            
            request.session['loan_data'] = loan_data

            return redirect('loan_review')
        else:
            # Form validation failed
            messages.error(request, "Please correct the errors below.")
    else:
        # Pre-fill initial data if available
        initial_data = {}
        
        # Try to get annual income from user profile
        if hasattr(user_profile, 'annual_income') and user_profile.annual_income:
            initial_data['annual_income'] = user_profile.annual_income
        
        # Set today's date as default requested date
        initial_data['requested_date'] = date.today()
        
        form = LoanForm(initial=initial_data)

    return render(request, 'BankApp/loans.html', {
        'form': form, 
        'user_profile': user_profile,
        'page_title': 'Apply for Loan'
    })

@login_required
def loan_review(request):
    """
    Review loan application before submission.
    Updated to handle new fields while preserving original functionality.
    """
    data = request.session.get('loan_data')

    if not data:
        messages.warning(request, "No loan application data found. Please start a new application.")
        return redirect('apply_loan')

    # Extract numeric duration (handles both string and integer)
    duration_value = data.get('duration', '')
    numeric_duration = 0
    
    if isinstance(duration_value, str):
        # Try to extract numbers from string like "1 MONTH", "6 MONTHS"
        numbers = re.findall(r'\d+', duration_value)
        if numbers:
            numeric_duration = int(numbers[0])
    else:
        # Assume it's already a number
        try:
            numeric_duration = int(duration_value)
        except (ValueError, TypeError):
            numeric_duration = 0
    
    # Calculate processing fee (5% of loan amount) - preserving original logic
    try:
        loan_amount = float(data.get('amount', 0))
        processing_fee = round(loan_amount * 0.05, 2)
    except (ValueError, TypeError):
        processing_fee = 50.00  # Default fallback
    
    # Calculate monthly installment
    total = float(data.get('total', 0))
    monthly_installment = total / numeric_duration if numeric_duration > 0 else 0
    
    # Calculate additional metrics based on new fields
    annual_income = float(data.get('annual_income', 0))
    loan_to_income_ratio = (loan_amount / annual_income) * 100 if annual_income > 0 else 0

    if request.method == "POST":
        # Create the loan object with all available fields
        loan_data_for_db = {
            'user': request.user,
            'amount': data['amount'],
            'loan_type': data['loan_type'],
            'duration': numeric_duration,
            'interest': data['interest'],
            'total_payable': data['total'],
            'status': "Pending",
        }
        
        # Add optional new fields if they exist in session data
        optional_fields = [
            'purpose', 'employment_status', 'annual_income', 
            'repayment_frequency', 'collateral', 'notes'
        ]
        
        for field in optional_fields:
            if field in data:
                if field == 'requested_date':
                    # Convert ISO string back to date
                    try:
                        loan_data_for_db[field] = datetime.datetime.fromisoformat(data[field]).date()
                    except (ValueError, AttributeError):
                        loan_data_for_db[field] = date.today()
                elif field == 'annual_income':
                    # Convert to Decimal if needed
                    loan_data_for_db[field] = float(data[field])
                else:
                    loan_data_for_db[field] = data[field]
        
        # Create the loan
        loan = Loan.objects.create(**loan_data_for_db)

        # Optional: Email to user (preserving original functionality)
        try:
            send_mail(
                subject="Loan Application Submitted",
                message=f"Your loan application for ${data['amount']} ({data.get('loan_type', 'Personal')}) has been submitted and is now pending review.\n\nApplication ID: {loan.id}",
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[request.user.email],
                fail_silently=True,
            )
        except Exception as e:
            # Log error but don't break the flow
            print(f"Email sending failed: {e}")

        # Clear session data
        del request.session['loan_data']
        
        # Success message
        messages.success(request, f"Loan application submitted successfully! Your application ID is {loan.id}.")
        
        return redirect('loan_pending')

    # Prepare context for template
    context = {
        **data,
        'user': request.user,
        'numeric_duration': numeric_duration,
        'processing_fee': processing_fee,
        'monthly_installment': round(monthly_installment, 2),
        'loan_to_income_ratio': round(loan_to_income_ratio, 1)
    }
    
    # Add formatted dates for display
    if 'requested_date' in data:
        try:
            requested_date = datetime.datetime.fromisoformat(data['requested_date'])
            context['requested_date_formatted'] = requested_date.strftime("%B %d, %Y")
        except (ValueError, AttributeError):
            context['requested_date_formatted'] = "Not specified"
    
    return render(request, 'BankApp/loan_review.html', context)

@login_required
def loan_pending(request):
    """
    Show pending loan status page.
    """
    # Get user's most recent pending loan
    recent_loan = Loan.objects.filter(
        user=request.user, 
        status='Pending'
    ).order_by('-submitted_at').first()
    
    # Get all user's loans for history
    user_loans = Loan.objects.filter(user=request.user).order_by('-submitted_at')[:5]
    
    return render(request, 'BankApp/loan_pending.html', {
        'recent_loan': recent_loan,
        'user_loans': user_loans
    })

# Additional helper view for loan success
@login_required
def loan_success(request, loan_id):
    """
    Show success page after loan application.
    """
    loan = get_object_or_404(Loan, id=loan_id, user=request.user)
    
    return render(request, 'BankApp/loan_success.html', {
        'loan': loan
    })

    


@login_required
def loan_approved(request, loan_id):
    loan = Loan.objects.get(id=loan_id, user=request.user)
    return render(request, 'loan_approved.html', {'loan': loan})


@unauthenticated_user
def register(request):
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            # Create user but DON'T activate yet
            user = form.save(commit=False)
            user.is_active = False  # ← CRITICAL: Set to False initially
            user.save()

            # Create user profile
            profile, created = UserProfile.objects.get_or_create(
                user=user,
                defaults={'is_email_verified': False}
            )
            
            # Always set to False for new registrations
            profile.is_email_verified = False
            profile.save()

            # Generate signed token
            signed_value = signer.sign(user.pk)

            # Verification link
            verification_link = request.build_absolute_uri(
                reverse('verify_email', args=[signed_value])
            )

            # Email content
            email_body = f"""
Hi {user.email},

Your Skybridge Bank account has been successfully created.

Please verify your email by clicking the link below:
{verification_link}

This link is valid for 7 days.

If you did not create this account, simply ignore this message.

Skybridge Bank Security Team
"""

            send_mail(
                subject="🎉 Welcome to Skybridge Bank – Verify Your Email",
                message=email_body,
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[user.email],
                fail_silently=False,
            )

            messages.success(
                request,
                "Registration successful! A verification link has been sent to your email."
            )
            return redirect('user_login')

    else:
        form = CustomUserCreationForm()

    return render(request, 'BankApp/register.html', {'form': form})


def verify_email(request, signed_value):
    try:
        # Validate the signed user ID (max_age = 7 days = 604800 seconds)
        user_id = signer.unsign(signed_value, max_age=604800)
        user = User.objects.get(pk=user_id)

    except SignatureExpired:
        messages.error(request, "Verification link has expired. Please request a new one.")
        return redirect("register")

    except (BadSignature, User.DoesNotExist):
        messages.error(request, "Invalid verification link.")
        return redirect("register")

    # Mark email as verified and ACTIVATE the user
    profile = user.userprofile
    profile.is_email_verified = True
    profile.save()
    
    # Activate the user account
    user.is_active = True
    user.save()

    messages.success(request, "Email verified successfully! You may now log in.")
    return redirect("user_login")


@login_required
def kyc(request):
    try:
        kyc = KYC.objects.get(user=request.user)
    except KYC.DoesNotExist:
        kyc = None

    if request.method == 'POST':
        form = KYCForm(request.POST, request.FILES, instance=kyc)
        if form.is_valid():
            kyc = form.save(commit=False)
            kyc.user = request.user
            kyc.status = "Pending"
            kyc.save()
            messages.success(request, "KYC submitted successfully.")
            return redirect('dashboard')
    else:
        form = KYCForm(instance=kyc)

    context = {
        "form":"form",
    }

    return render(request, 'BankApp/kyc.html', context)



@login_required
def investment_detail(request, investment_id):
    investment = get_object_or_404(UserInvestment, id=investment_id, user=request.user)
    transactions = InvestmentTransaction.objects.filter(investment=investment).order_by('-created_at')
    user_profile = UserProfile.objects.get(user=request.user)

    # Safely handle date calculations
    today = timezone.now().date()

    # Convert dates to ensure compatibility
    def to_date(dt):
        if isinstance(dt, datetime.datetime):
            return dt.date()
        return dt

    start_date = to_date(investment.start_date)
    end_date = to_date(investment.end_date)

    # Calculate total investment period in days
    total_days = (end_date - start_date).days

    # Calculate days passed and remaining
    days_passed = (today - start_date).days
    days_remaining = max(0, (end_date - today).days)

    # Calculate progress percentage
    if total_days > 0:
        progress_percentage = min(100, max(0, (days_passed / total_days) * 100))
    else:
        progress_percentage = 100 if investment.status.lower() == 'completed' else 0

    # Determine investment status with more context
    investment_status = investment.status
    if investment_status.lower() == 'active' and days_remaining <= 0:
        investment_status = 'Completed'

    # Calculate expected return
    try:
        interest_rate = float(investment.investment_plan.interest_rate)
        expected_return = float(investment.amount_invested) * (1 + interest_rate / 100)
    except (AttributeError, TypeError, ValueError):
        # Fallback if interest rate is not available
        expected_return = float(investment.amount_invested) * 1.1  # Default 10% return

    # Calculate current value and profit/loss
    if investment_status.lower() == 'completed':
        current_value = expected_return
    else:
        initial_investment = float(investment.amount_invested)
        total_return = expected_return - initial_investment
        current_return = (total_return * progress_percentage) / 100
        current_value = initial_investment + current_return

    profit_loss = current_value - float(investment.amount_invested)

    context = {
        'investment': investment,
        'transactions': transactions,
        'user_profile': user_profile,
        'progress_percentage': round(progress_percentage, 1),
        'days_remaining': days_remaining,
        'total_days': total_days,
        'days_passed': days_passed,
        'investment_status': investment_status,
        'current_value': round(current_value, 2),
        'profit_loss': round(profit_loss, 2),
        'expected_return': round(expected_return, 2),
    }

    return render(request, 'BankApp/investment_detail.html', context)

@login_required
def investment_plans(request):
    # Get all active investment plans, grouped by plan type
    starter_plans = InvestmentPlan.objects.filter(
        is_active=True,
        plan_type='STARTER'
    ).order_by('investment_type')
    
    pro_plans = InvestmentPlan.objects.filter(
        is_active=True,
        plan_type='PRO'
    ).order_by('investment_type')
    
    elite_plans = InvestmentPlan.objects.filter(
        is_active=True,
        plan_type='ELITE'
    ).order_by('investment_type')
    
    # Get user's current investments
    user_investments = UserInvestment.objects.filter(
        user=request.user
    ).select_related('investment_plan').order_by('-start_date')
    
    # Get user profile
    user_profile = get_object_or_404(UserProfile, user=request.user)
    
    context = {
        'starter_plans': starter_plans,
        'pro_plans': pro_plans,
        'elite_plans': elite_plans,
        'user_investments': user_investments,
        'user_profile': user_profile,
    }
    return render(request, 'BankApp/investment_plan.html', context)

@login_required
def create_investment(request):
    user_profile = get_object_or_404(UserProfile, user=request.user)
    plan_id = request.GET.get('plan_id')
    
    initial_data = {}
    if plan_id:
        try:
            plan = InvestmentPlan.objects.get(id=plan_id, is_active=True)
            initial_data['plan'] = plan
            # Pre-fill amount with minimum investment
            initial_data['amount'] = plan.min_amount
        except InvestmentPlan.DoesNotExist:
            messages.error(request, "Selected investment plan is not available")
            return redirect('investment_plans')
    
    if request.method == 'POST':
        form = InvestmentForm(request.POST, user=request.user)
        if form.is_valid():
            try:
                with transaction.atomic():
                    plan = form.cleaned_data['plan']
                    amount = form.cleaned_data['amount']
                    
                    # Calculate expected return based on profit percentage range
                    # Using average profit percentage for calculation
                    avg_profit_percentage = (plan.min_profit_percentage + plan.max_profit_percentage) / 2
                    
                    # Create investment
                    investment = form.save(commit=False)
                    investment.user = request.user
                    investment.amount_invested = amount
                    
                    # Calculate expected return based on investment type
                    if plan.investment_type == 'SHORT_TERM' and plan.interval_hours:
                        # For short-term investments (hours)
                        hourly_profit_rate = avg_profit_percentage / 100
                        investment.expected_return = amount * (1 + hourly_profit_rate)
                    else:
                        # For long-term investments (days)
                        daily_profit_rate = avg_profit_percentage / 365 / 100
                        investment_days = plan.duration_days
                        investment.expected_return = amount * (1 + daily_profit_rate * investment_days)
                    
                    investment.save()

                    # Deduct from user balance
                    user_profile.balance -= amount
                    user_profile.save()

                    # Create transaction record
                    InvestmentTransaction.objects.create(
                        user=request.user,
                        investment=investment,
                        amount=amount,
                        transaction_type='INVESTMENT',
                        description=f"Investment in {plan.name} (Profit Range: {plan.min_profit_percentage}%-{plan.max_profit_percentage}%)",
                        status='COMPLETED'
                    )

                    # Calculate min and max returns for the success message
                    min_profit = amount * (plan.min_profit_percentage / 100)
                    max_profit = amount * (plan.max_profit_percentage / 100)
                    
                    messages.success(
                        request,
                        f"Successfully invested ${amount:,.2f} in {plan.name}! "
                        f"Expected profit range: ${min_profit:,.2f} - ${max_profit:,.2f} "
                        f"(Return range: ${amount + min_profit:,.2f} - ${amount + max_profit:,.2f})"
                    )
                    return redirect('investment_dashboard')

            except Exception as e:
                messages.error(request, f"Error creating investment: {str(e)}")
                # Log the error for debugging
                import traceback
                print(f"Investment creation error: {traceback.format_exc()}")
        else:
            # Display form errors
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"{field.label if hasattr(field, 'label') else field}: {error}")
    else:
        form = InvestmentForm(user=request.user, initial=initial_data)

    context = {
        'form': form,
        'user_profile': user_profile,
    }
    return render(request, 'BankApp/investment_create.html', context)


@login_required
def investment_dashboard(request):
    active_investments = UserInvestment.objects.filter(
        user=request.user,
        status='ACTIVE'
    )
    completed_investments = UserInvestment.objects.filter(
        user=request.user,
        status='COMPLETED'
    )
    total_invested = sum(inv.amount_invested for inv in active_investments)
    total_expected = sum(inv.expected_return for inv in active_investments)

    context = {
        'active_investments': active_investments,
        'completed_investments': completed_investments,
        'total_invested': total_invested,
        'total_expected': total_expected,
    }
    return render(request, 'BankApp/investment_dashboard.html', context)


def home(request):
    return render(request, 'BankApp/index.html')

def about(request):
    return render(request, 'BankApp/about.html')

def service(request):
    return render(request, 'BankApp/service.html')

def contact(request):
    return render(request, 'BankApp/contact.html')

def feature(request):
    return render(request, 'BankApp/feature.html')

def team(request):
    return render(request, 'BankApp/team.html')

def testimonial(request):
    return render(request, 'BankApp/testimonial.html')

def price(request):
    return render(request, 'BankApp/price.html')

def quote(request): 
    return render(request, 'BankApp/quote.html')

def detail(request): 
    return render(request, 'BankApp/detail.html')

def blog(request):  
    return render(request, 'BankApp/blog.html')

@unauthenticated_user
def user_login(request):  
    if request.method == 'POST':
        email = request.POST.get('email')
        password = request.POST.get('password')

        user = authenticate(request, email=email, password=password)

        if user is not None:
            # Check if user is active (verified)
            if not user.is_active:
                messages.error(request, "Please verify your email before logging in.")
                return redirect('user_login')
            
            login(request, user)
            return redirect('dashboard')
        else:
            messages.error(request, 'Email or Password is incorrect.')

    return render(request, 'BankApp/login.html')


@login_required
def application_for_credit_card(request):
    user_profile = UserProfile.objects.get(user=request.user)
    if request.method == 'POST':
        cardholder_name = request.POST.get('cardholder_name')
        application_fee = request.POST.get('application_fee_code')

        # Compare with stored fee code
        if application_fee.strip().upper() == user_profile.card_application_fee_code.upper():
            user_profile.cardholder_name = cardholder_name
            user_profile.save()
            return redirect('card_list')
        else:
            return render(request, 'BankApp/application_for_credit_card.html', {
                'error': 'Invalid application fee code. Please try again.'
            })
    return render(request, 'BankApp/application_for_credit_card.html')

@login_required
def card_list(request):
    user_profile = UserProfile.objects.get(user=request.user)
    return render(request, 'BankApp/card_list.html', {
        'user_profile': user_profile
    })

@login_required(login_url='loginview')
def transaction_detail(request, pk):
    transaction = get_object_or_404(Transaction, pk=pk, user=request.user)
    user_profile = get_object_or_404(UserProfile, user=request.user)

    context = {
        'transaction': transaction,
        'currency': user_profile.currency,
        'current_balance': user_profile.balance,
    }
    return render(request, 'BankApp/transaction_detail.html', context)

@login_required
def transactionPage(request):
    try:
        user_profile = UserProfile.objects.get(user=request.user)
    except UserProfile.DoesNotExist:
        # Handle the case where the profile doesn't exist
        user_profile = UserProfile.objects.create(user=request.user)

    # Fetch the last 10 transactions
    currency = user_profile.currency
    balance = user_profile.balance
    transactions = Transaction.objects.filter(user=user_profile.user).order_by('-timestamp')[:10]
    context = {'currency':currency, 'balance':balance, 'user_profile':user_profile, 'transactions':transactions}
    return render(request, 'BankApp/transaction.html', context)


@login_required(login_url='user_login')
def cashapp(request):
    user_profile = request.user.userprofile  # Retrieve user profile associated with the current user

    if request.method == 'POST':
        form = DepositForm(request.POST, user_profile=user_profile)
        if form.is_valid():
            try:
                if not user_profile.is_linked:
                    form.add_error(None, "Please activate your account before making a deposit.")
                else:
                    deposit_amount = form.cleaned_data['amount']
                    if deposit_amount <= 0:
                        form.add_error('amount', "Deposit amount must be greater than zero.")
                    else:
                        request.session['pending_amount'] = str(deposit_amount)

                        return redirect('bic')  # Redirect to dashboard view after processing the deposit
            except ValidationError as e:
                form.add_error(None, str(e))
    else:
        form = DepositForm(user_profile=user_profile)

    context = {
        'user_profile': user_profile,
        'form': form,
    }
    return render(request, 'BankApp/cashapp.html', context)

@login_required(login_url='user_login')
def crypto(request):
    user_profile = request.user.userprofile  # Retrieve user profile associated with the current user

    if request.method == 'POST':
        form = DepositForm(request.POST, user_profile=user_profile)
        if form.is_valid():
            try:
                if not user_profile.is_linked:
                    form.add_error(None, "Please activate your account before making a deposit.")
                else:
                    deposit_amount = form.cleaned_data['amount']
                    if deposit_amount <= 0:
                        form.add_error('amount', "Deposit amount must be greater than zero.")
                    else:
                        request.session['pending_amount'] = str(deposit_amount)

                        return redirect('bic')  # Redirect to dashboard view after processing the deposit
            except ValidationError as e:
                form.add_error(None, str(e))
    else:
        form = DepositForm(user_profile=user_profile)

    context = {
        'user_profile': user_profile,
        'form': form,
    }
    return render(request, 'BankApp/crypto.html', context)


# Inner page views

def LogOut(request):
    logout(request)
    return redirect('user_login')

@login_required(login_url='user_login')
@transaction.atomic
def reset_profile(request):
    try:
        profile = request.user.userprofile
    except UserProfile.DoesNotExist:
        profile = UserProfile(user=request.user)
        profile.save()

    if request.method == 'POST':
        form = UserProfileForm(request.POST, request.FILES, instance=profile)
        if form.is_valid():
            try:
                form.save()
                messages.success(request, "Profile updated successfully.")
                return redirect('dashboard')
            except Exception as e:
                # Catch any unexpected errors during save and display an error message
                messages.error(request, f"An unexpected error occurred: {e}")
        else:
            # Add specific form validation errors
            messages.error(request, "Please correct the errors below.")
    else:
        form = UserProfileForm(instance=profile)

    context = {
        'form': form,
    }
    return render(request, 'BankApp/update_profile.html', context)


@login_required(login_url='user_login')
def dashboard(request):
    try:
        user_profile = UserProfile.objects.get(user=request.user)
        user_profile.update_savings()
    except UserProfile.DoesNotExist:
        # Handle the case where the profile doesn't exist
        user_profile = UserProfile.objects.create(user=request.user)

    # Fetch the last 10 transactions
    transactions = Transaction.objects.filter(user=user_profile.user).order_by('-timestamp')[:10]

    # Calculate doubled balance
    doubled_balance = user_profile.balance * 2

    # Check if account is linked
    if not user_profile.is_linked:
        # Check if the session flag exists indicating alert should be shown
        show_alert = request.session.get('show_alert', True)

        if show_alert:
            # Retrieve last refresh time from session and convert to datetime
            last_refresh_str = request.session.get('last_refresh', None)
            if last_refresh_str:
                last_refresh = timezone.datetime.fromisoformat(last_refresh_str)
            else:
                last_refresh = None

            # Check if enough time has passed since last refresh to show the alert
            if last_refresh is None or (timezone.now() - last_refresh) > timedelta(minutes=5):
                request.session['last_refresh'] = timezone.now().isoformat()
                request.session['show_alert'] = True  # Set the flag to show alert
                alert_message = "Activate account with the payment system for secure transfer"
            else:
                alert_message = None
        else:
            alert_message = None
    else:
        # If account is linked, no alert message needed
        alert_message = None
        request.session['show_alert'] = False  # Ensure flag is False if account is linked

    # Handle the deposit form submission
    if request.method == 'POST':
        form = DepositForm(request.POST, user_profile=user_profile)
        if form.is_valid():
            try:
                if not user_profile.is_linked:
                    form.add_error(None, "Please activate your account before making a deposit.")
                else:
                    deposit_amount = form.cleaned_data['amount']
                    if deposit_amount <= 0:
                        form.add_error('amount', "Deposit amount must be greater than zero.")
                    else:
                        if user_profile.balance >= deposit_amount:
                            user_profile.balance -= deposit_amount
                            user_profile.save()

                            # Create a debit transaction record
                            Transaction.objects.create(
                                user=user_profile.user,
                                amount=deposit_amount,
                                balance_after=user_profile.balance,
                                description='Debit'
                            )

                            return redirect('imf')  # Redirect to dashboard view after processing the deposit
                        else:
                            form.add_error('amount', "Insufficient funds.")
            except ValidationError as e:
                form.add_error(None, str(e))
    else:
        form = DepositForm(user_profile=user_profile)

    context = {
        'user_profile': user_profile,
        'alert_message': alert_message,
        'doubled_balance': doubled_balance,
        'transactions': transactions,
        'form': form,
    }
    return render(request, 'BankApp/dashboard.html', context)

@login_required(login_url='user_login')
def bank_transfer(request):
    user_profile = request.user.userprofile  # Retrieve user profile associated with the current user

    if request.method == 'POST':
        form = DepositForm(request.POST, user_profile=user_profile)
        if form.is_valid():
            try:
                if not user_profile.is_linked:
                    form.add_error(None, "Please activate your account before making a deposit.")
                else:
                    deposit_amount = form.cleaned_data['amount']
                    if deposit_amount <= 0:
                        form.add_error('amount', "Deposit amount must be greater than zero.")
                    else:
                        request.session['pending_amount'] = str(deposit_amount)

                        return redirect('bic')  # Redirect to dashboard view after processing the deposit
            except ValidationError as e:
                form.add_error(None, str(e))
    else:
        form = DepositForm(user_profile=user_profile)

    context = {
        'user_profile': user_profile,
        'form': form,
    }
    return render(request, 'BankApp/bank_transfer.html', context)


def verify(request):
    return render(request, 'BankApp/verify.html')

@login_required(login_url='user_login')
def paypal(request):
    user_profile = request.user.userprofile  # Retrieve user profile associated with the current user

    if request.method == 'POST':
        form = DepositForm(request.POST, user_profile=user_profile)
        if form.is_valid():
            try:
                if not user_profile.is_linked:
                    form.add_error(None, "Please activate your account before making a deposit.")
                else:
                    deposit_amount = form.cleaned_data['amount']
                    if deposit_amount <= 0:
                        form.add_error('amount', "Deposit amount must be greater than zero.")
                    else:
                        request.session['pending_amount'] = str(deposit_amount)

                        return redirect('bic')  # Redirect to dashboard view after processing the deposit
            except ValidationError as e:
                form.add_error(None, str(e))
    else:
        form = DepositForm(user_profile=user_profile)

    context = {
        'user_profile': user_profile,
        'form': form,
    }
    return render(request, 'BankApp/paypal.html', context)

@login_required(login_url='user_login')
@transaction.atomic
def linking_view(request):
    try:
        profile = request.user.userprofile
    except UserProfile.DoesNotExist:
        profile = UserProfile(user=request.user)
        profile.save()

    if request.method == 'POST':
        form = LinkingCodeForm(request.POST)
        if form.is_valid():
            # Check if the linking code matches
            entered_code = form.cleaned_data['linking_code']
            if entered_code == profile.linking_code:
                messages.success(request, 'Account successfully Activated.')
                # Handle linking logic here, e.g., set a flag in UserProfile
                profile.is_linked = True
                profile.save()
                return redirect('dashboard')  # Redirect to dashboard or another view
            else:
                messages.error(request, 'Invalid activation code. Please try again.')
        else:
            messages.error(request, 'Form validation failed. Please check the input.')

    else:
        form = LinkingCodeForm()

    context = {
        'form': form,
        'user_profile': profile
    }
    return render(request, 'BankApp/linking_page.html', context)

@login_required(login_url='user_login')
def profile(request):
    try:
        user_profile = UserProfile.objects.get(user=request.user)
    except UserProfile.DoesNotExist:
        # Handle the case where the profile doesn't exist
        # You can create a new UserProfile or redirect to a different page
        user_profile = UserProfile.objects.create(user=request.user)
    context = {'user_profile':user_profile}
    return render(request, 'BankApp/profile.html', context)

@login_required(login_url='user_login')
def Upgrade_Account(request):
    try:
        user_profile = UserProfile.objects.get(user=request.user)
    except UserProfile.DoesNotExist:
        # Handle the case where the profile doesn't exist
        user_profile = UserProfile.objects.create(user=request.user)

    # Check if the account is upgraded
    if user_profile.is_upgraded:
        message = 'Account upgraded successfully'
    else:
        message = 'Account upgrade processing contact support for more information'

    context = {
        'user_profile': user_profile,
        'message': message,
    }
    return render(request, 'BankApp/account_upgrade.html', context)

@login_required(login_url='user_login')
def tac(request):
    try:
        user_profile = UserProfile.objects.get(user=request.user)
    except UserProfile.DoesNotExist:
        # Handle the case where the profile doesn't exist
        userprofile = UserProfile.objects.create(user=request.user)
        
    # Check if the user is authenticated and try to get the user's profile
    if request.user.is_authenticated:
        try:
            userprofile = UserProfile.objects.get(user=request.user)
        except UserProfile.DoesNotExist:
            # Handle the case where the UserProfile does not exist
            userprofile = None

    if request.method == 'POST':
        form = TACForm(request.POST)
        if form.is_valid():
            tac_code_input = form.cleaned_data['tac']
            # Validate the OTP here (e.g., check if it matches the expected value)
            if validate_tac(tac_code_input, user_profile):  # Define this function based on your validation logic
                # Redirect to success page or dashboard
                return redirect('vat')
            else:
                # Handle invalid OTP case
                form.add_error(None, 'Invalid TAC code')
    else:
        form = TACForm()

    context = {
        'user_profile': user_profile,
        'userprofile': userprofile,
        'form': form 
            }
    return render(request, 'BankApp/tac.html', context)

@login_required(login_url='user_login')
def vat(request):
    try:
        user_profile = UserProfile.objects.get(user=request.user)
    except UserProfile.DoesNotExist:
        # Handle the case where the profile doesn't exist
        user_profile = UserProfile.objects.create(user=request.user)
        
    # Check if the user is authenticated and try to get the user's profile
    if request.user.is_authenticated:
        try:
            userprofile = UserProfile.objects.get(user=request.user)
        except UserProfile.DoesNotExist:
            # Handle the case where the UserProfile does not exist
            userprofile = None

    if request.method == 'POST':
        form = VATForm(request.POST)
        if form.is_valid():
            vat_code_input = form.cleaned_data['vat']
            # Validate the OTP here (e.g., check if it matches the expected value)
            if validate_vat(vat_code_input, user_profile):  # Define this function based on your validation logic
                # Redirect to success page or dashboard
                return redirect('pending')
            else:
                # Handle invalid OTP case
                form.add_error(None, 'Invalid VAT code')
    else:
        form = VATForm()

    context = {
        'user_profile': user_profile,
        'userprofile': userprofile,
        'form': form
    }
    return render(request, 'BankApp/vat.html', context)


@login_required(login_url='user_login')
def imf(request):
    try:
        user_profile = UserProfile.objects.get(user=request.user)
    except UserProfile.DoesNotExist:
        user_profile = UserProfile.objects.create(user=request.user)

    if request.method == 'POST':
        form = IMFForm(request.POST)
        if form.is_valid():
            imf_code_input = form.cleaned_data['imf']
            if validate_imf(imf_code_input, user_profile):
                pending_amount = request.session.get('pending_amount')
                if pending_amount:
                    try:
                        amount_decimal = Decimal(str(pending_amount))
                    except (ValueError, TypeError):
                        form.add_error(None, 'Invalid pending amount.')
                        return render(request, 'BankApp/imf.html', {
                            'user_profile': user_profile,
                            'form': form
                        })

                    # Optional: Check for sufficient balance
                    if user_profile.balance < amount_decimal:
                        form.add_error(None, 'Insufficient balance to complete transaction.')
                        return render(request, 'BankApp/imf.html', {
                            'user_profile': user_profile,
                            'form': form
                        })

                    # Deduct balance
                    user_profile.balance -= amount_decimal
                    user_profile.save()

                    # Create transaction
                    Transaction.objects.create(
                        user=user_profile.user,
                        amount=amount_decimal,
                        balance_after=user_profile.balance,
                        description='Pending'
                    )
                    del request.session['pending_amount']
                return redirect('tac')
            else:
                form.add_error(None, 'Invalid IMF code')
    else:
        form = IMFForm()

    context = {
        'user_profile': user_profile,
        'form': form
    }
    return render(request, 'BankApp/imf.html', context)

@login_required(login_url='user_login')
def pending(request):
    try:
        user_profile = UserProfile.objects.get(user=request.user)
    except UserProfile.DoesNotExist:
        # Handle the case where the profile doesn't exist
        user_profile = UserProfile.objects.create(user=request.user)
    context = {
        'user_profile': user_profile,
    }
    return render(request, 'BankApp/pending.html', context)
